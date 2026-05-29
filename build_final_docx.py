#!/usr/bin/env python3
"""生成阿坝州医共体方案最终合订本Word文档 — 全宋体专业排版"""

import re, os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ============================================================
# 配置
# ============================================================
OUTPUT_FILE = "E:/claude/output/word/阿坝州医共体方案-合订本-最终版.docx"
FONT = "宋体"
C_DARK = RGBColor(0x1A, 0x3C, 0x6E)
C_ACCENT = RGBColor(0x2B, 0x57, 0x9A)
C_BODY = RGBColor(0x22, 0x22, 0x22)
C_GRAY = RGBColor(0x66, 0x66, 0x66)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_TABLE_BG = "1A3C6E"
C_ALT_ROW = "F2F6FA"

# ============================================================
# 工具函数
# ============================================================

def new_para(doc, text="", font_size=Pt(11), bold=False, color=C_BODY, align=None,
             before=0, after=4, line=1.5, font_name=None, indent=None):
    """添加格式化段落"""
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        fn = font_name or FONT
        run.font.name = fn
        run._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
        run.font.size = font_size
        run.bold = bold
        run.font.color.rgb = color
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if indent:
        pf.left_indent = Cm(indent)
    return p

def heading(doc, text, level):
    """添加样式化标题"""
    sizes = {0: Pt(24), 1: Pt(20), 2: Pt(15), 3: Pt(12.5), 4: Pt(11.5)}
    befores = {0: 36, 1: 22, 2: 16, 3: 12, 4: 8}
    afters = {0: 14, 1: 10, 2: 7, 3: 5, 4: 4}
    p = new_para(doc, text, sizes.get(level, Pt(12)), True, C_DARK,
                 before=befores.get(level, 10), after=afters.get(level, 5))
    if level <= 1:
        pPr = p._p.get_or_add_pPr()
        sz = "12" if level == 0 else "6" if level == 1 else "3"
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="{sz}" w:space="4" w:color="1A3C6E"/></w:pBdr>')
        pPr.append(pBdr)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """添加格式化表格"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = FONT; r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        r.font.size = Pt(9); r.bold = True; r.font.color.rgb = C_WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shade(cell, C_TABLE_BG)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ci >= len(headers): break
            cell = table.rows[ri+1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = FONT; r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
            r.font.size = Pt(8.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if ri % 2 == 1:
                set_cell_shade(cell, C_ALT_ROW)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def set_cell_shade(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def page_break(doc):
    doc.add_page_break()

# ============================================================
# 封面
# ============================================================
def build_cover(doc):
    for _ in range(8): new_para(doc, before=0, after=0, line=1.0)

    new_para(doc, "阿坝州紧密型县域医共体", Pt(26), True, C_DARK, WD_ALIGN_PARAGRAPH.CENTER, after=2, line=1.2)
    new_para(doc, "信息化建设项目", Pt(26), True, C_DARK, WD_ALIGN_PARAGRAPH.CENTER, after=6, line=1.2)
    new_para(doc, "建设方案（合订本·最终版）", Pt(14), False, C_ACCENT, WD_ALIGN_PARAGRAPH.CENTER, after=24, line=1.5)

    # 装饰线
    p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="8" w:color="1A3C6E"/></w:pBdr>')
    pPr.append(pBdr)

    for _ in range(6): new_para(doc, before=0, after=0, line=1.0)

    info = [
        ("项目总投资", "9,600万元（已到位3,000万元）"),
        ("建设范围", "阿坝州13县（市）全域"),
        ("方案风格", "问题导向型（痛点驱动+ROI量化+利旧优先）"),
        ("编制日期", "2026年5月"),
        ("版 本 号", "V2.0 最终版"),
        ("密    级", "内部"),
    ]
    for label, value in info:
        p = new_para(doc, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2, line=1.8)
        r1 = p.add_run(f"{label}：")
        r1.font.name = FONT; r1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        r1.font.size = Pt(11); r1.bold = True; r1.font.color.rgb = C_DARK
        r2 = p.add_run(value)
        r2.font.name = FONT; r2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        r2.font.size = Pt(11); r2.font.color.rgb = C_BODY

    page_break(doc)

# ============================================================
# 目录页
# ============================================================
def build_toc(doc):
    heading(doc, "目  录", 0)
    toc_items = [
        ("第一章", "阿坝医共体的五大生死线", 3),
        ("第二章", "根因分析与破局思路", 4),
        ("第三章", "基础能力层：先止血（技术架构+平台+远程医疗）", 6),
        ("第四章", "业务提升层：再造血（PACS/体检/导诊/院感+系统集成）", 10),
        ("第五章", "三医联动闭环设计（医保+药品+绩效）", 14),
        ("第六章", "AI轻量化赋能策略（5场景）", 16),
        ("第七章", "运营造血机制（增收+控费+效率）", 18),
        ("第八章", "投入产出全景测算（投资明细+ROI分析）", 20),
        ("第九章", "实施与风控（详细计划+培训+Go-Live+风险）", 23),
        ("附录一", "全国医共体信息化建设案例分析（7案例+适配矩阵）", 27),
        ("附录二", "逐系统功能规格说明书（8大系统）", 30),
        ("附录三", "AI模型训练技术规范（影像+NLP+MLOps+高原特色）", 34),
        ("附录四", "存量信息化家底普查报告（13县46系统评估）", 38),
        ("附录五", "查漏补缺清单与政策依据手册（16模块）", 40),
        ("附录六", "数据质量校验报告（55项逐项校验）", 43),
    ]
    for num, title, page in toc_items:
        p = new_para(doc, before=2, after=2, line=1.8)
        r1 = p.add_run(f"{num}  ")
        r1.font.name = FONT; r1._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        r1.font.size = Pt(11); r1.bold = True; r1.font.color.rgb = C_DARK
        r2 = p.add_run(title)
        r2.font.name = FONT; r2._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
        r2.font.size = Pt(11); r2.font.color.rgb = C_BODY

    page_break(doc)

# ============================================================
# 核心指标摘要页
# ============================================================
def build_exec_summary(doc):
    heading(doc, "核心指标摘要", 0)

    heading(doc, "一、项目基本指标", 2)
    add_table(doc,
        ["指标", "数值", "说明"],
        [["项目总投资", "9,600万元", "中央+省+地方配套"],
         ["已到位资金", "3,000万元", "一期启动资金"],
         ["建设周期", "36个月（三期）", "每期12个月"],
         ["覆盖范围", "13县（市）+171乡镇+1,168村卫生室", "全域覆盖"],
         ["建设策略", "查漏补缺·避免重复·不大拆大建·适配改造优先·性价比突出", "五项铁律"]])

    heading(doc, "二、投资结构", 2)
    add_table(doc,
        ["投资类别", "金额(万元)", "占比", "核心模块"],
        [["基础平台", "3,350", "34.9%", "州级信息平台+13县平台+基础设施"],
         ["业务应用", "2,995", "31.2%", "PACS+体检+导诊+院感疾控"],
         ["三医联动", "1,300", "13.5%", "医保AI审核+药品供应链+绩效"],
         ["AI赋能", "770", "8.0%", "影像AI+医保AI+慢病AI+传染病AI+处方AI"],
         ["配套服务", "1,185", "12.3%", "培训+运维+集成+预备金"]])

    heading(doc, "三、核心回报指标", 2)
    add_table(doc,
        ["指标", "建设前", "建设后(目标)", "依据"],
        [["州域内就诊率", "79.69%", "≥85%", "三明/东台/滨海经验"],
         ["检查检验互认率", "≈0%", "≥70%", "国卫办63号文要求"],
         ["DRG审核覆盖率", "<5%(人工)", "100%(AI)", "梧州/尤溪模式"],
         ["处方合格率", "≈85%(估)", "≥95%", "滨海85%→99.86%"],
         ["影像报告时效", "数小时-天", "<30分钟(常规)", "重庆/安徽模式"],
         ["静态回收期", "—", "≈4.2年", "含2年建设期"],
         ["年可量化回报(成熟期)", "—", "≈4,000万元", "增收+控费+效率"]])

    heading(doc, "四、利旧指标", 2)
    add_table(doc,
        ["资源类别", "利旧率目标", "策略"],
        [["软件系统", "保留48%+改造39%=87%", "保留主流品牌、改造接口、仅替换极少数"],
         ["硬件设备", "≥60%", "服务器/存储/网络设备/终端最大化复用"],
         ["网络设施", "≥80%", "电子政务外网已覆盖大部分乡镇"],
         ["历史数据", "100%", "全部诊疗数据迁移至新平台"]])

    page_break(doc)

# ============================================================
# 主文：风格B方案整合
# ============================================================
def build_main_body(doc):
    """从风格B MD文件读取并转换核心内容到Word"""
    md_path = "E:/claude/output/md/阿坝州医共体方案-B-问题导向型.md"

    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # 按 ## 分割章节
    sections = re.split(r'\n(?=## )', content)

    for sec in sections:
        lines = sec.strip().split('\n')
        if not lines: continue

        first = lines[0].strip()
        if first.startswith('# ') and not first.startswith('## '):
            heading(doc, first[2:].strip(), 0)
        elif first.startswith('## '):
            heading(doc, first[3:].strip(), 1)
        elif first.startswith('### '):
            heading(doc, first[4:].strip(), 2)
        elif first.startswith('#### '):
            heading(doc, first[5:].strip(), 3)

        # 处理内容
        in_code = False
        in_table = False
        table_lines = []

        for line in lines[1:]:
            stripped = line.strip()

            if stripped.startswith('```'):
                in_code = not in_code
                continue
            if in_code: continue

            if stripped.startswith('|') and stripped.endswith('|'):
                if '---' in stripped:
                    if table_lines:
                        # 解析表
                        headers = [c.strip() for c in table_lines[0].split('|')[1:-1]]
                        rows = []
                        for tl in table_lines[1:]:
                            cells = [c.strip() for c in tl.split('|')[1:-1]]
                            rows.append(cells)
                        add_table(doc, headers, rows)
                    table_lines = []
                else:
                    table_lines.append(stripped)
                continue

            if table_lines:
                for tl in table_lines:
                    new_para(doc, tl, Pt(9.5), color=C_GRAY)
                table_lines = []
                continue

            if stripped == '':
                new_para(doc, before=0, after=0, line=0.5)
            elif stripped.startswith('> '):
                new_para(doc, stripped[2:], Pt(10), color=C_GRAY, indent=0.8)
            elif stripped.startswith('- ') or stripped.startswith('* '):
                text = re.sub(r'^[-*]\s+', '', stripped)
                new_para(doc, f"  • {text}", Pt(10.5), indent=0.5)
            elif stripped.startswith('**') and stripped.endswith('**') and len(stripped) < 80:
                new_para(doc, stripped[2:-2], Pt(11), True, C_ACCENT, before=8, after=3)
            elif not stripped.startswith('#') and not stripped.startswith('---'):
                # 处理行内粗体
                if '**' in stripped and not stripped.startswith('|'):
                    p = doc.add_paragraph()
                    pf = p.paragraph_format
                    pf.space_before = Pt(1); pf.space_after = Pt(3); pf.line_spacing = 1.5
                    parts = re.split(r'(\*\*[^*]+\*\*)', stripped)
                    for part in parts:
                        r = p.add_run(part[2:-2] if part.startswith('**') else part)
                        r.font.name = FONT; r._element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
                        r.font.size = Pt(10.5)
                        if part.startswith('**'): r.bold = True
                else:
                    new_para(doc, stripped, Pt(10.5))

# ============================================================
# 附录整合
# ============================================================
def build_appendix_placeholder(doc):
    """为图表JPG预留位置说明"""
    heading(doc, "图表索引", 2)
    charts_list = [
        "图1: 总体架构图 — 1+13+N三层架构",
        "图2: 技术架构图 — 六层技术栈",
        "图3: 数据架构图 — 数据全生命周期",
        "图7: 远程影像AI辅助诊断业务流程",
        "图8: 医保AI三级漏斗审核流程",
        "图13: 八统一管理逻辑图",
        "图14: 三医联动逻辑图",
        "图16: 投资结构分解与优先级矩阵",
        "图19: 各建设模块投入产出对比(ROI热力图)",
    ]
    for i, ch in enumerate(charts_list):
        jpg_name = f"图{i+1:02d}_{ch.split('—')[0].strip().replace(':', '_')}.jpg"
        new_para(doc, f"{ch}", Pt(9.5), True, C_ACCENT)
        new_para(doc, f"  → 对应JPG文件: output/charts/jpg/{jpg_name}", Pt(8.5), color=C_GRAY)
        new_para(doc, "  [ 图表预留位 — 请将对应JPG文件插入此处 ]", Pt(8.5), color=C_GRAY, indent=1)
        new_para(doc, before=0, after=6)

# ============================================================
# 主流程
# ============================================================
def main():
    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.8); section.right_margin = Cm(2.8)
        section.page_width = Cm(21.0); section.page_height = Cm(29.7)

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = FONT
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5

    # 构建各部分
    build_cover(doc)
    build_toc(doc)
    build_exec_summary(doc)
    build_main_body(doc)

    # 图表预留
    page_break(doc)
    build_appendix_placeholder(doc)

    # 保存
    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)
    size_kb = os.path.getsize(OUTPUT_FILE) / 1024
    print(f"合订本已生成: {OUTPUT_FILE} ({size_kb:.0f}KB)")

if __name__ == '__main__':
    main()
