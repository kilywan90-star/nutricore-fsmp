#!/usr/bin/env python3
"""将方案编制任务书MD转换为美观的Word文档（全宋体）"""

import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy

# ============================================================
# 配置
# ============================================================
INPUT_FILE = "E:/claude/output/调研笔记/方案编制任务书-最终版.md"
OUTPUT_FILE = "E:/claude/output/word/阿坝州医共体方案编制任务书.docx"

FONT_NAME = "宋体"
FONT_SIZE_BODY = Pt(11)
FONT_SIZE_H1 = Pt(22)
FONT_SIZE_H2 = Pt(16)
FONT_SIZE_H3 = Pt(13)
FONT_SIZE_H4 = Pt(11.5)
FONT_SIZE_TABLE = Pt(9.5)
FONT_SIZE_CODE = Pt(8.5)
FONT_SIZE_COVER_TITLE = Pt(28)
FONT_SIZE_COVER_SUB = Pt(14)

COLOR_PRIMARY = RGBColor(0x1A, 0x3C, 0x6E)      # 深蓝
COLOR_ACCENT = RGBColor(0x2B, 0x57, 0x9A)        # 中蓝
COLOR_HEADING_BG = RGBColor(0xE8, 0xEF, 0xF5)     # 浅蓝背景
COLOR_TABLE_HEADER = RGBColor(0x1A, 0x3C, 0x6E)   # 表头深蓝
COLOR_TABLE_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)  # 表头白字
COLOR_TABLE_ROW_ALT = RGBColor(0xF2, 0xF6, 0xFA)  # 交替行浅蓝
COLOR_BORDER = RGBColor(0x8C, 0xA3, 0xC2)         # 边框蓝灰
COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)
COLOR_DARK = RGBColor(0x22, 0x22, 0x22)

# ============================================================
# 工具函数
# ============================================================

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val","single")}" '
            f'w:sz="{val.get("sz","4")}" '
            f'w:color="{val.get("color","8CA3C2")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line

def add_formatted_paragraph(doc, text, font_size=None, bold=False, color=None,
                             alignment=None, spacing_before=0, spacing_after=0,
                             font_name=None, line_spacing=None):
    """添加格式化段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name or FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name or FONT_NAME)
    run.font.size = font_size or FONT_SIZE_BODY
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if alignment is not None:
        p.alignment = alignment
    set_paragraph_spacing(p, spacing_before, spacing_after, line_spacing)
    return p

def style_table(table, header_rows=1):
    """美化表格：交替行颜色、边框、表头"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 设置整个表格字体
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.name = FONT_NAME
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                    run.font.size = FONT_SIZE_TABLE
            # 交替行背景
            if row_idx >= header_rows and row_idx % 2 == 1:
                set_cell_shading(cell, "F2F6FA")
            # 表头
            if row_idx < header_rows:
                set_cell_shading(cell, "1A3C6E")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True

def add_styled_heading(doc, text, level, font_size, color):
    """添加样式化标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = font_size
    run.bold = True
    run.font.color.rgb = color

    # 标题前间距
    before_map = {1: 24, 2: 18, 3: 14, 4: 10}
    after_map = {1: 10, 2: 8, 3: 6, 4: 4}
    set_paragraph_spacing(p, before_map.get(level, 8), after_map.get(level, 4))

    # H1底部装饰线
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="8" w:space="4" w:color="1A3C6E"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    elif level == 2:
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="4" w:space="2" w:color="8CA3C2"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)

    return p

def add_body_text(doc, text):
    """添加正文段落"""
    if not text.strip():
        return
    p = doc.add_paragraph()
    run = p.add_run(text.strip())
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_BODY
    run.font.color.rgb = COLOR_DARK
    set_paragraph_spacing(p, 1, 4, 1.5)

def add_code_block(doc, lines):
    """添加代码块"""
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = FONT_SIZE_CODE
        run.font.color.rgb = COLOR_GRAY
        set_paragraph_spacing(p, 0, 0, 1.0)
        p.paragraph_format.left_indent = Cm(0.5)

def add_bullet(doc, text, indent_level=0):
    """添加项目符号"""
    p = doc.add_paragraph()
    prefix = '    ' * indent_level
    run = p.add_run(f"{prefix}• {text.strip()}")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_BODY
    set_paragraph_spacing(p, 1, 2, 1.4)
    p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.8)

def add_description_item(doc, term, description):
    """添加描述列表项"""
    p = doc.add_paragraph()
    run_term = p.add_run(f"{term}：")
    run_term.font.name = FONT_NAME
    run_term._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run_term.font.size = FONT_SIZE_BODY
    run_term.bold = True
    run_desc = p.add_run(description)
    run_desc.font.name = FONT_NAME
    run_desc._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run_desc.font.size = FONT_SIZE_BODY
    set_paragraph_spacing(p, 1, 2, 1.4)

def add_separator(doc):
    """添加分隔线"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="2" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    set_paragraph_spacing(p, 12, 6)


# ============================================================
# 封面
# ============================================================

def create_cover(doc):
    """创建封面页"""
    # 空行填充
    for _ in range(6):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, 0, 0, 1.0)

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("阿坝州紧密型县域医共体")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_COVER_TITLE
    run.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    set_paragraph_spacing(p, 0, 4, 1.2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("信息化建设项目")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_COVER_TITLE
    run.bold = True
    run.font.color.rgb = COLOR_PRIMARY
    set_paragraph_spacing(p, 0, 8, 1.2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("方案编制任务书")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = FONT_SIZE_COVER_SUB
    run.font.color.rgb = COLOR_ACCENT
    set_paragraph_spacing(p, 0, 30, 1.5)

    # 分隔装饰线
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="12" w:space="8" w:color="1A3C6E"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # 副标题信息
    for _ in range(4):
        doc.add_paragraph()

    cover_info = [
        ("编制日期", "2026年5月"),
        ("版   本   号", "V2.0（最终版）"),
        ("编 制 单 位", "阿坝州卫生健康委员会（拟）"),
        ("密       级", "内部"),
    ]
    for label, value in cover_info:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_l = p.add_run(f"{label}：")
        run_l.font.name = FONT_NAME
        run_l._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        run_l.font.size = Pt(11)
        run_l.bold = True
        run_l.font.color.rgb = COLOR_PRIMARY
        run_v = p.add_run(value)
        run_v.font.name = FONT_NAME
        run_v._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
        run_v.font.size = Pt(11)
        run_v.font.color.rgb = COLOR_DARK
        set_paragraph_spacing(p, 2, 2, 1.8)

    # 封面后分页
    doc.add_page_break()


# ============================================================
# 正文解析与生成
# ============================================================

def parse_and_generate(doc, md_path):
    """解析Markdown并生成Word内容"""
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []
    in_list = False
    list_buffer = []
    skip_until_content = True  # 跳过文件开头的YAML frontmatter

    while i < len(lines):
        line = lines[i]
        raw = line.rstrip('\n')

        # 代码块处理
        if raw.strip().startswith('```'):
            if in_code_block:
                if code_buffer:
                    add_code_block(doc, code_buffer)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(raw)
            i += 1
            continue

        # 跳过frontmatter
        if skip_until_content:
            if raw.strip() == '---' and i > 0:
                skip_until_content = False
            i += 1
            continue
        if raw.strip() == '---':
            i += 1
            continue

        stripped = raw.strip()

        # 表格处理
        if stripped.startswith('|') and stripped.endswith('|'):
            if '---' in stripped:
                # 表格分隔行
                if table_buffer:
                    # 解析表头
                    headers = [c.strip() for c in table_buffer[0].split('|')[1:-1]]
                    # 创建表格
                    num_rows = len(table_buffer)
                    table = doc.add_table(rows=num_rows, cols=len(headers))
                    table.style = 'Table Grid'

                    # 填充表头
                    for col_idx, header in enumerate(headers):
                        cell = table.rows[0].cells[col_idx]
                        cell.text = ''
                        p = cell.paragraphs[0]
                        run = p.add_run(header)
                        run.font.name = FONT_NAME
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                        run.font.size = FONT_SIZE_TABLE
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    # 填充数据行
                    for row_idx, data_row in enumerate(table_buffer[1:], 1):
                        cells = [c.strip() for c in data_row.split('|')[1:-1]]
                        for col_idx, cell_text in enumerate(cells):
                            if col_idx < len(headers):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = ''
                                p = cell.paragraphs[0]
                                run = p.add_run(cell_text)
                                run.font.name = FONT_NAME
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                                run.font.size = FONT_SIZE_TABLE

                    style_table(table)
                    doc.add_paragraph()  # 表后空行
                table_buffer = []
            else:
                table_buffer.append(raw)
            i += 1
            continue

        # 表格结束，刷新缓冲区
        if table_buffer and not stripped.startswith('|'):
            # 把缓冲内容当正文处理
            for buf_line in table_buffer:
                add_body_text(doc, buf_line)
            table_buffer = []
            # 继续处理当前行

        # 标题
        if stripped.startswith('# ') and not stripped.startswith('## '):
            add_styled_heading(doc, stripped[2:], 1, FONT_SIZE_H1, COLOR_PRIMARY)
        elif stripped.startswith('## '):
            add_styled_heading(doc, stripped[3:], 2, FONT_SIZE_H2, COLOR_PRIMARY)
        elif stripped.startswith('### '):
            add_styled_heading(doc, stripped[4:], 3, FONT_SIZE_H3, COLOR_ACCENT)
        elif stripped.startswith('#### '):
            add_styled_heading(doc, stripped[5:], 4, FONT_SIZE_H4, COLOR_DARK)
        elif stripped.startswith('**') and stripped.endswith('**') and len(stripped) < 80:
            # 加粗短行 = 小标题
            add_formatted_paragraph(doc, stripped[2:-2], FONT_SIZE_BODY, bold=True,
                                     color=COLOR_PRIMARY, spacing_before=8, spacing_after=4)
        elif stripped.startswith('> '):
            # 引用
            text = stripped[2:]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = FONT_NAME
            run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
            run.font.size = FONT_SIZE_BODY
            run.italic = True
            run.font.color.rgb = COLOR_GRAY
            set_paragraph_spacing(p, 2, 2, 1.4)
            p.paragraph_format.left_indent = Cm(0.8)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = re.sub(r'^[-*]\s+', '', stripped)
            add_bullet(doc, text)
        elif stripped.startswith('  - ') or stripped.startswith('    - '):
            text = re.sub(r'^\s+[-*]\s+', '', stripped)
            add_bullet(doc, text, 1)
        elif stripped.startswith('|---'):
            # 表格分隔行（不在表缓冲中时忽略）
            pass
        elif stripped == '':
            if not in_list:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, 0, 0, 0.5)
        elif stripped in ['---', '***', '___']:
            add_separator(doc)
        elif stripped.startswith('!['):
            # 图片（跳过或占位）
            pass
        else:
            # 处理行内格式
            text = stripped
            # 粗体 **text**
            if '**' in text and not (text.startswith('|') and text.endswith('|')):
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.name = FONT_NAME
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
                    run.font.size = FONT_SIZE_BODY
                set_paragraph_spacing(p, 1, 3, 1.5)
            elif text.startswith('|') and text.endswith('|'):
                # 单行表格
                pass
            else:
                add_body_text(doc, text)

        i += 1

    # 刷新残留的表格缓冲
    if table_buffer:
        for buf_line in table_buffer:
            add_body_text(doc, buf_line)


# ============================================================
# 页面设置
# ============================================================

def setup_page(doc):
    """设置页面格式"""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    style.font.size = FONT_SIZE_BODY
    style.paragraph_format.line_spacing = 1.5


# ============================================================
# 主流程
# ============================================================

def main():
    doc = Document()

    # 页面设置
    setup_page(doc)

    # 封面
    create_cover(doc)

    # 正文内容
    parse_and_generate(doc, INPUT_FILE)

    # 保存
    os.makedirs(os.path.dirname(OUTPUT_FILE), exist_ok=True)
    doc.save(OUTPUT_FILE)
    print(f"Word文档已生成: {OUTPUT_FILE}")

    # 文件大小
    size_kb = os.path.getsize(OUTPUT_FILE) / 1024
    print(f"文件大小: {size_kb:.0f} KB")


if __name__ == '__main__':
    main()
