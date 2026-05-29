"""
标书导出 API — 支持导出为 Word (.docx) 格式
"""
import re
import logging
from io import BytesIO

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from db.sqlite import get_db, Project

logger = logging.getLogger(__name__)
router = APIRouter()


class ExportRequest(BaseModel):
    project_id: int
    format: str = "docx"
    include_toc: bool = True
    split_sections: bool = False


def _set_run_font(run, cn_name: str):
    """给 run 设置中文字体"""
    fm = {"宋体": "SimSun", "黑体": "SimHei", "楷体": "KaiTi"}
    en = fm.get(cn_name, "SimSun")
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_name)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)


def md_to_docx(md_content: str, title: str = "") -> BytesIO:
    """Markdown → Word 文档"""
    doc = Document()

    # A4 页面
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.18)
    sec.right_margin = Cm(3.18)

    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line[2:].strip())
            run.font.size = Pt(22)
            run.font.bold = True
            _set_run_font(run, '黑体')

        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line[3:].strip())
            run.font.size = Pt(16)
            run.font.bold = True
            _set_run_font(run, '黑体')

        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(line[4:].strip())
            run.font.size = Pt(14)
            run.font.bold = True
            _set_run_font(run, '黑体')

        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('—' * 40)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(180, 180, 180)

        elif line.strip().startswith('|') and line.strip().endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1

            rows_data = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip().split('|')[1:-1]]
                if all(re.match(r'^[-:\s]+$', c) for c in cells):
                    continue
                if cells:
                    rows_data.append(cells)

            if rows_data:
                table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
                table.style = 'Table Grid'
                for r, row_data in enumerate(rows_data):
                    for c, cell_text in enumerate(row_data):
                        if c < len(table.rows[r].cells):
                            cell = table.rows[r].cells[c]
                            cell.text = cell_text
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    run.font.size = Pt(10)
                doc.add_paragraph()

        elif '**' in line:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.font.bold = True
                else:
                    p.add_run(part)

        elif line.strip():
            p = doc.add_paragraph(line.strip())

        i += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── API ──────────────────────────────────────────

@router.post("/docx")
async def export_docx(req: ExportRequest, db: Session = Depends(get_db)):
    """导出标书为 Word (.docx) 文件"""
    project = db.query(Project).filter(Project.id == req.project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    if not project.bid_content:
        raise HTTPException(status_code=400, detail="标书内容为空")

    try:
        import traceback as _tb
        buf = md_to_docx(project.bid_content, project.name or "标书")
        from urllib.parse import quote
        safe_name = re.sub(r'[\\/*?:"<>|]', '', project.name or "标书")
        encoded_name = quote(f"{safe_name}_投标文件.docx")
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"}
        )
    except Exception as e:
        _tb.print_exc()
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")


@router.post("/pdf")
async def export_pdf(req: ExportRequest, db: Session = Depends(get_db)):
    """导出标书为 HTML（可在浏览器中打开并打印为 PDF）"""
    project = db.query(Project).filter(Project.id == req.project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="项目不存在")
    if not project.bid_content:
        raise HTTPException(status_code=400, detail="标书内容为空")

    import markdown as md_lib
    html_body = md_lib.markdown(project.bid_content, extensions=['tables', 'fenced_code', 'toc'])
    html = f"""<!DOCTYPE html>
<html lang=\"zh-CN\">
<head><meta charset=\"utf-8\"><title>{project.name or '标书'}</title>
<style>
@page {{ size: A4; margin: 2.54cm 3.18cm; }}
body {{ font-family: \"SimSun\",serif; font-size: 12pt; line-height: 1.8; }}
h1 {{ font-family: \"SimHei\",sans-serif; font-size: 22pt; text-align: center; }}
h2 {{ font-family: \"SimHei\",sans-serif; font-size: 16pt; }}
h3 {{ font-family: \"SimHei\",sans-serif; font-size: 14pt; }}
table {{ border-collapse: collapse; width: 100%; }}
td, th {{ border: 1px solid #000; padding: 6px; font-size: 10pt; }}
</style></head><body>{html_body}</body></html>"""

    buf = BytesIO(html.encode('utf-8'))
    from urllib.parse import quote
    safe_name = re.sub(r'[\\/*?:"<>|]', '', project.name or "标书")
    encoded_name = quote(f"{safe_name}_投标文件.html")
    return StreamingResponse(
        buf,
        media_type="text/html; charset=utf-8",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_name}"}
    )
