import os
from pypdf import PdfReader
from docx import Document
import pandas as pd
def extract_text_from_file(file_path: str, file_type: str) -> str:
    """
    从不同类型文件中提取文本内容
    :param file_path: 文件路径
    :param file_type: 文件类型（pdf/docx/xlsx等）
    :return: 提取的文本内容
    """
    file_type = file_type.lower()
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type in ["doc", "docx"]:
        return extract_text_from_docx(file_path)
    elif file_type in ["xls", "xlsx"]:
        return extract_text_from_excel(file_path)
    elif file_type == "txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"不支持的文件类型: {file_type}")
def extract_text_from_pdf(file_path: str) -> str:
    """从PDF文件提取文本"""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
        return text
    except Exception as e:
        raise Exception(f"PDF解析失败: {str(e)}")
def extract_text_from_docx(file_path: str) -> str:
    """从Word文档提取文本"""
    try:
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join([cell.text.strip() for cell in row.cells])
                if row_text.strip():
                    text += row_text + "\n"
            text += "\n"
        return text
    except Exception as e:
        raise Exception(f"Word文档解析失败: {str(e)}")
def extract_text_from_excel(file_path: str) -> str:
    """从Excel文件提取文本"""
    try:
        text = ""
        excel_file = pd.ExcelFile(file_path)
        for sheet_name in excel_file.sheet_names:
            text += f"=== Sheet: {sheet_name} ===\n"
            df = pd.read_excel(excel_file, sheet_name=sheet_name)
            # 将DataFrame转为文本
            text += df.to_string(index=False, na_rep="") + "\n\n"
        return text
    except Exception as e:
        raise Exception(f"Excel文件解析失败: {str(e)}")
def extract_text_from_txt(file_path: str) -> str:
    """从文本文件提取内容"""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        # 尝试用gbk编码读取
        try:
            with open(file_path, "r", encoding="gbk") as f:
                return f.read()
        except Exception as e:
            raise Exception(f"文本文件读取失败: {str(e)}")
