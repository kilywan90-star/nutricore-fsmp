"""将MD格式方案文件转换为Word(.docx)和TXT格式"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

BASE_DIR = "E:/claude/output"
MD_DIR = os.path.join(BASE_DIR, "md")
WORD_DIR = os.path.join(BASE_DIR, "word")
TXT_DIR = os.path.join(BASE_DIR, "txt")

FILES = [
    "阿坝州医共体方案-A-标准合规型.md",
    "阿坝州医共体方案-B-问题导向型.md",
    "阿坝州医共体方案-C-创新引领型.md",
]

def md_to_docx(md_path, docx_path):
    """Convert Markdown to Word document with basic formatting"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    in_table = False

    for line in lines:
        # Code blocks
        if line.strip().startswith('```'):
            in_code_block = not in_code_block
            continue

        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            continue

        # Skip horizontal rules
        if line.strip() == '---':
            continue

        # Headings
        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            for run in p.runs:
                run.font.name = '黑体'
        elif line.startswith('## '):
            title = line[3:].strip()
            p = doc.add_heading(title, level=1)
            for run in p.runs:
                run.font.name = '黑体'
        elif line.startswith('### '):
            title = line[4:].strip()
            p = doc.add_heading(title, level=2)
            for run in p.runs:
                run.font.name = '黑体'
        elif line.startswith('#### '):
            title = line[5:].strip()
            p = doc.add_heading(title, level=3)
            for run in p.runs:
                run.font.name = '黑体'

        # Tables (simple detection for | --- | format)
        elif line.startswith('|') and line.strip().endswith('|'):
            if '---' in line:
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            # For simplicity, add as formatted paragraph
            p = doc.add_paragraph()
            for i, cell in enumerate(cells):
                if i > 0:
                    p.add_run('  |  ').font.size = Pt(10)
                run = p.add_run(cell)
                run.font.size = Pt(10)
                run.font.name = '宋体'

        # Bullet points
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = re.sub(r'^[-*]\s+', '', line.strip())
            p = doc.add_paragraph(text, style='List Bullet')

        # Blockquotes
        elif line.startswith('> '):
            text = line[2:].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)

        # Bold text markers
        elif line.strip().startswith('**') and line.strip().endswith('**'):
            text = line.strip()[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True

        # Empty lines
        elif line.strip() == '':
            doc.add_paragraph()

        # Regular text
        else:
            text = line.strip()
            if text:
                # Handle inline bold **text**
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*[^*]+\*\*)', text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = p.add_run(part)
            else:
                doc.add_paragraph()

    doc.save(docx_path)
    print(f"Word saved: {docx_path}")


def md_to_txt(md_path, txt_path):
    """Strip markdown formatting and save as plain text"""
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Strip basic markdown formatting
    content = re.sub(r'```[\s\S]*?```', '', content)  # Code blocks
    content = re.sub(r'`([^`]+)`', r'\1', content)     # Inline code
    content = re.sub(r'\*\*([^*]+)\*\*', r'\1', content)  # Bold
    content = re.sub(r'\|', ' ', content)               # Table pipes
    content = re.sub(r'---+\n', '', content)             # HR
    content = re.sub(r'^\s*>+\s?', '', content, flags=re.MULTILINE)  # Blockquotes
    content = re.sub(r'^[-*]\s+', '  • ', content, flags=re.MULTILINE)  # Bullets
    content = re.sub(r'^#+\s+', '', content, flags=re.MULTILINE)  # Headings
    content = re.sub(r'\n{3,}', '\n\n', content)  # Collapse blank lines

    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f"TXT saved: {txt_path}")


if __name__ == '__main__':
    os.makedirs(WORD_DIR, exist_ok=True)
    os.makedirs(TXT_DIR, exist_ok=True)

    for filename in FILES:
        md_path = os.path.join(MD_DIR, filename)
        base_name = os.path.splitext(filename)[0]

        docx_path = os.path.join(WORD_DIR, base_name + '.docx')
        txt_path = os.path.join(TXT_DIR, base_name + '.txt')

        print(f"Converting: {filename}")
        md_to_docx(md_path, docx_path)
        md_to_txt(md_path, txt_path)

    print("\nAll conversions complete!")
