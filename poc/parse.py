"""
文档解析模块：PDF/Word/图片 → Markdown 文本
依赖: pip install pymupdf python-docx pillow
OCR 可选: pip install paddleocr (扫描件必须)
"""

import os
import sys
from pathlib import Path

def parse_pdf(filepath: str) -> str:
    """解析 PDF，返回 Markdown 文本。优先用 PyMuPDF 提取文本，失败则走 OCR。"""
    import fitz  # PyMuPDF

    doc = fitz.open(filepath)
    texts = []
    for page in doc:
        text = page.get_text(sort=True)
        if text.strip():
            texts.append(text)
        else:
            # 纯图片页，走 OCR
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")
            texts.append(_ocr_image_bytes(img_bytes))

    result = "\n\n".join(texts)
    if not result.strip():
        raise ValueError(f"未能从 PDF 提取任何文本: {filepath}")
    return result


def parse_docx(filepath: str) -> str:
    """解析 Word 文档，提取段落和表格。"""
    from docx import Document

    doc = Document(filepath)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            style = para.style.name if para.style else ""
            if "Heading" in style or "heading" in style or "标题" in style:
                parts.append(f"## {para.text.strip()}")
            else:
                parts.append(para.text.strip())

    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if rows:
            # 插入表头分隔线
            header_sep = "|" + "|".join(["---"] * len(table.rows[0].cells)) + "|"
            rows.insert(1, header_sep)
            parts.append("\n".join(rows))

    result = "\n\n".join(parts)
    if not result.strip():
        raise ValueError(f"未能从 Word 提取任何文本: {filepath}")
    return result


def parse_image(filepath: str) -> str:
    """OCR 识别图片中的文字。"""
    from PIL import Image

    img = Image.open(filepath)
    return _ocr_image(img)


def parse_file(filepath: str) -> dict:
    """
    入口函数：根据文件扩展名自动选择解析器。
    返回 {"text": str, "file_type": str, "page_count": int}
    """
    ext = Path(filepath).suffix.lower()
    path = str(filepath)

    parsers = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".doc": parse_docx,  # 需要预先转为 docx
        ".png": parse_image,
        ".jpg": parse_image,
        ".jpeg": parse_image,
        ".bmp": parse_image,
        ".tiff": parse_image,
    }

    parser = parsers.get(ext)
    if not parser:
        raise ValueError(f"不支持的文件格式: {ext}")

    text = parser(path)
    return {
        "text": text,
        "file_type": ext,
        "filename": Path(filepath).name,
        "char_count": len(text),
    }


def _ocr_image_bytes(img_bytes: bytes) -> str:
    """对图片字节进行 OCR。需要 PaddleOCR。"""
    try:
        from paddleocr import PaddleOCR
        import numpy as np
        from PIL import Image
        import io

        ocr = PaddleOCR(lang="ch", show_log=False)
        img = Image.open(io.BytesIO(img_bytes))
        img_np = np.array(img)
        results = ocr.ocr(img_np)
        if not results or not results[0]:
            return ""
        lines = []
        for line in results[0]:
            text = line[1][0]
            lines.append(text)
        return "\n".join(lines)
    except ImportError:
        print("⚠ PaddleOCR 未安装，返回空文本。安装: pip install paddleocr", file=sys.stderr)
        return ""
    except Exception as e:
        print(f"⚠ OCR 失败: {e}", file=sys.stderr)
        return ""


def _ocr_image(img) -> str:
    """对 PIL Image 进行 OCR。需要 PaddleOCR。"""
    try:
        from paddleocr import PaddleOCR
        import numpy as np

        ocr = PaddleOCR(lang="ch", show_log=False)
        img_np = np.array(img)
        results = ocr.ocr(img_np)
        if not results or not results[0]:
            return ""
        return "\n".join(line[1][0] for line in results[0])
    except ImportError:
        print("⚠ PaddleOCR 未安装。安装: pip install paddleocr", file=sys.stderr)
        return ""
    except Exception as e:
        print(f"⚠ OCR 失败: {e}", file=sys.stderr)
        return ""


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(f"用法: python {sys.argv[0]} <文件路径>")
        sys.exit(1)

    result = parse_file(sys.argv[1])
    print(f"文件: {result['filename']}")
    print(f"类型: {result['file_type']}")
    print(f"字符数: {result['char_count']}")
    print(f"\n{'='*60}\n")
    print(result["text"][:2000])
